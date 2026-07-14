from collections import Counter
from datetime import date, datetime, time, timedelta
from pathlib import Path
import os

from django.conf import settings
from django.utils import timezone
from docx import Document
from docx.shared import Inches
import requests

from apps.ai_results.models import Alert

from .models import MonitorReport

DEEPSEEK_API_URL = os.getenv("DEEPSEEK_API_URL", "https://api.deepseek.com/chat/completions")
DEEPSEEK_MODEL = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")

REPORT_PERIODS = {
    "00:00-06:00": (0, 6),
    "06:00-12:00": (6, 12),
    "12:00-18:00": (12, 18),
    "18:00-24:00": (18, 24),
}


def default_report_date():
    return timezone.localdate()


def previous_completed_period(reference=None):
    """Return the last fully completed 6-hour local reporting period."""
    local_now = timezone.localtime(reference or timezone.now())
    end_hour = (local_now.hour // 6) * 6
    if end_hour == 0:
        start_date = local_now.date() - timedelta(days=1)
        start_hour, end_date = 18, local_now.date()
    else:
        start_date = local_now.date()
        start_hour, end_date = end_hour - 6, local_now.date()

    tz = timezone.get_current_timezone()
    start = timezone.make_aware(datetime.combine(start_date, time(start_hour, 0)), tz)
    end = timezone.make_aware(datetime.combine(end_date, time(end_hour, 0)), tz)
    return start, end, _period_label(start, end)


def report_period(report_date: date, period_label: str | None = None, end_at_now: bool = False):
    """Resolve a report period.

    With a period label, report_date means the date on which the period starts.
    Without a label, the latest completed 6-hour period is used.
    """
    if period_label:
        return period_for_label(report_date, period_label)
    if end_at_now:
        start, end, label = current_partial_period()
        return start, end, label
    return previous_completed_period()


def period_for_label(report_date: date, period_label: str):
    if period_label not in REPORT_PERIODS:
        raise ValueError("periodLabel must be one of: " + ", ".join(REPORT_PERIODS))
    start_hour, end_hour = REPORT_PERIODS[period_label]
    tz = timezone.get_current_timezone()
    start = timezone.make_aware(datetime.combine(report_date, time(start_hour, 0)), tz)
    end_date = report_date + timedelta(days=1) if end_hour == 24 else report_date
    end_time = time(0, 0) if end_hour == 24 else time(end_hour, 0)
    end = timezone.make_aware(datetime.combine(end_date, end_time), tz)
    return start, end, period_label


def current_partial_period(reference=None):
    local_now = timezone.localtime(reference or timezone.now())
    start_hour = (local_now.hour // 6) * 6
    start_date = local_now.date()
    tz = timezone.get_current_timezone()
    start = timezone.make_aware(datetime.combine(start_date, time(start_hour, 0)), tz)
    end = timezone.now()
    return start, end, _period_label(start, _period_end_for_start(start))


def generate_latest_period_report() -> MonitorReport:
    period_start, period_end, label = previous_completed_period()
    return generate_monitor_report(
        report_date=timezone.localtime(period_start).date(),
        period_start=period_start,
        period_end=period_end,
        period_label=label,
    )


def generate_monitor_report(
    report_date: date | None = None,
    period_label: str | None = None,
    end_at_now: bool = False,
    period_start=None,
    period_end=None,
) -> MonitorReport:
    if period_start is None or period_end is None:
        target_date = report_date or default_report_date()
        period_start, period_end, resolved_label = report_period(
            target_date,
            period_label=period_label,
            end_at_now=end_at_now,
        )
    else:
        resolved_label = period_label or _period_label(period_start, period_end)
        target_date = report_date or timezone.localtime(period_start).date()

    target_date = report_date or timezone.localtime(period_start).date()
    report, _ = MonitorReport.objects.update_or_create(
        period_start=period_start,
        period_end=period_end,
        defaults={
            "report_date": target_date,
            "period_label": resolved_label,
            "status": MonitorReport.Status.GENERATING,
            "error_message": "",
        },
    )

    alerts = list(
        Alert.objects.select_related("camera", "event")
        .filter(occurred_at__gte=period_start, occurred_at__lt=period_end)
        .order_by("occurred_at", "id")
    )

    try:
        alert_items = [_alert_item(alert) for alert in alerts]
        ai_summary = _generate_management_advice(period_start, period_end, alert_items)
        content = _build_report_content(period_start, period_end, resolved_label, alert_items, ai_summary)
        document_path = _write_word_document(period_start, period_end, resolved_label, content, alert_items, ai_summary)

        report.alert_count = len(alerts)
        report.high_alert_count = sum(1 for alert in alerts if alert.level == "high")
        report.pending_alert_count = sum(1 for alert in alerts if alert.status == Alert.Status.PENDING)
        report.ai_summary = ai_summary
        report.content = content
        report.document_path = document_path
        report.status = MonitorReport.Status.GENERATED
        report.error_message = ""
        report.generated_at = timezone.now()
        report.save()
    except Exception as exc:
        report.status = MonitorReport.Status.FAILED
        report.error_message = str(exc)
        report.generated_at = timezone.now()
        report.save(update_fields=["status", "error_message", "generated_at", "updated_at"])

    return report


def _generate_management_advice(period_start, period_end, alert_items):
    api_key = os.getenv("DEEPSEEK_API_KEY", "").strip()
    fallback = _fallback_management_advice(alert_items)
    if not api_key:
        return fallback

    response = requests.post(
        DEEPSEEK_API_URL,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json={
            "model": DEEPSEEK_MODEL,
            "messages": [
                {
                    "role": "system",
                    "content": (
                        "你是工厂安全监控管理助手。请只基于提供的告警数据，"
                        "输出一到两句中文管理建议，不超过120字，不要编造未提供的信息。"
                    ),
                },
                {"role": "user", "content": _build_advice_prompt(period_start, period_end, alert_items)},
            ],
            "temperature": 0.2,
        },
        timeout=30,
    )
    response.raise_for_status()
    payload = response.json()
    content = payload.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
    return content[:500] if content else fallback


def _build_advice_prompt(period_start, period_end, alert_items):
    lines = [
        f"统计周期：{_fmt(period_start)} 至 {_fmt(period_end)}",
        f"告警总数：{len(alert_items)}",
        "请输出一到两句管理建议，聚焦高危、重复区域、重复类型和处置优先级。",
        "告警事件：",
    ]
    if not alert_items:
        lines.append("无告警事件。")
        return "\n".join(lines)
    for item in alert_items[:80]:
        lines.append(
            f"- 时间={item['occurredAt']}; 摄像头={item['cameraName']}; 类型={item['eventType']}; "
            f"等级={item['level']}; 状态={item['status']}; 描述={item['description'] or '无'}"
        )
    return "\n".join(lines)


def _fallback_management_advice(alert_items):
    if not alert_items:
        return "本时段未产生告警事件，建议保持现有巡检频率并继续关注重点区域。"
    level_counts = Counter(item["level"] for item in alert_items)
    type_counts = Counter(item["eventType"] for item in alert_items)
    top_type, top_count = type_counts.most_common(1)[0]
    high_count = level_counts.get("high", 0)
    return f"本时段共发生 {len(alert_items)} 起告警，其中高危 {high_count} 起，主要类型为 {top_type}（{top_count} 起）。建议优先复核高危和待处理事件，并针对高频类型加强现场巡检。"


def _build_report_content(period_start, period_end, period_label, alert_items, ai_summary):
    level_counts = Counter(item["level"] for item in alert_items)
    type_counts = Counter(item["eventType"] for item in alert_items)
    status_counts = Counter(item["status"] for item in alert_items)
    lines = [
        "FactoryVision AI 告警时段报告",
        "",
        f"统计周期：{_fmt(period_start)} 至 {_fmt(period_end)}",
        f"统计时段：{period_label}",
        "",
        "一、AI 管理建议",
        ai_summary,
        "",
        "二、告警统计",
        f"- 总告警数：{len(alert_items)}",
        f"- 高危告警：{level_counts.get('high', 0)}",
        f"- 中危告警：{level_counts.get('medium', 0)}",
        f"- 待处理告警：{status_counts.get(Alert.Status.PENDING, 0)}",
        "",
        "三、告警类型统计",
    ]
    lines.extend([f"- {key}: {value}" for key, value in type_counts.items()] or ["- 无"])
    lines.extend(["", "四、事件明细"])
    if not alert_items:
        lines.append("本时段无告警事件。")
        return "\n".join(lines)

    for index, item in enumerate(alert_items, start=1):
        lines.extend(
            [
                f"{index}. {item['title']}",
                f"   时间：{item['occurredAt']}",
                f"   摄像头：{item['cameraName']}",
                f"   类型：{item['eventType']}",
                f"   等级：{item['level']}",
                f"   状态：{item['status']}",
                f"   说明：{item['description'] or '无'}",
                f"   关键帧：{item['keyframeUrl'] or '暂无'}",
            ]
        )
    return "\n".join(lines)


def _write_word_document(period_start, period_end, period_label, content, alert_items, ai_summary):
    reports_dir = Path(settings.MEDIA_ROOT) / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)
    filename = f"monitor-report-{timezone.localtime(period_start).strftime('%Y%m%d-%H%M')}-{timezone.localtime(period_end).strftime('%H%M')}.docx"
    path = reports_dir / filename

    document = Document()
    document.add_heading("FactoryVision AI 告警时段报告", level=1)
    document.add_paragraph(f"统计周期：{_fmt(period_start)} 至 {_fmt(period_end)}")
    document.add_paragraph(f"统计时段：{period_label}")
    document.add_heading("一、AI 管理建议", level=2)
    document.add_paragraph(ai_summary)
    document.add_heading("二、事件明细", level=2)
    if not alert_items:
        document.add_paragraph("本时段无告警事件。")
    for index, item in enumerate(alert_items, start=1):
        document.add_heading(f"{index}. {item['title']}", level=3)
        document.add_paragraph(f"时间：{item['occurredAt']}")
        document.add_paragraph(f"摄像头：{item['cameraName']}")
        document.add_paragraph(f"类型：{item['eventType']}    等级：{item['level']}    状态：{item['status']}")
        document.add_paragraph(f"说明：{item['description'] or '无'}")
        image_path = item.get("keyframeFilePath")
        if image_path and Path(image_path).exists():
            document.add_paragraph("关键帧：")
            try:
                document.add_picture(str(image_path), width=Inches(5.8))
            except Exception:
                document.add_paragraph(f"关键帧文件无法嵌入：{item.get('keyframeUrl') or item.get('keyframePath')}")
        else:
            document.add_paragraph("关键帧：暂无")

    document.add_page_break()
    document.add_heading("文本摘要", level=2)
    for block in content.splitlines():
        text = block.strip()
        if text:
            document.add_paragraph(text)
    document.save(path)
    return f"reports/{filename}"


def _alert_item(alert):
    keyframe_path = _alert_keyframe_path(alert)
    return {
        "id": alert.id,
        "title": alert.title,
        "eventType": alert.event_type,
        "level": alert.level,
        "status": alert.status,
        "cameraId": alert.camera_id,
        "cameraName": alert.camera.name if alert.camera else "未关联摄像头",
        "occurredAt": timezone.localtime(alert.occurred_at).strftime("%Y-%m-%d %H:%M:%S"),
        "description": alert.description,
        "keyframePath": keyframe_path,
        "keyframeUrl": _media_url(keyframe_path),
        "keyframeFilePath": _media_file_path(keyframe_path),
    }


def _alert_keyframe_path(alert):
    candidates = [
        alert.snapshot_path,
        getattr(alert.event, "snapshot_path", ""),
    ]
    payload = getattr(alert.event, "payload", {}) or {}
    media = payload.get("media") if isinstance(payload, dict) else {}
    if isinstance(media, dict):
        candidates.extend([media.get("keyframePath"), media.get("keyframeUrl")])
    if isinstance(payload, dict):
        candidates.extend([payload.get("keyframePath"), payload.get("snapshotPath"), payload.get("snapshotUrl")])
    return next((str(value).strip() for value in candidates if value), "")


def _media_url(path):
    if not path:
        return ""
    if str(path).startswith(("http://", "https://", "/")):
        return str(path)
    return f"{settings.MEDIA_URL.rstrip('/')}/{str(path).lstrip('/')}"


def _media_file_path(path):
    if not path or str(path).startswith(("http://", "https://")):
        return ""
    relative = str(path).removeprefix(settings.MEDIA_URL).lstrip("/")
    return str(Path(settings.MEDIA_ROOT) / relative)


def _period_end_for_start(period_start):
    return period_start + timedelta(hours=6)


def _period_label(period_start, period_end):
    local_start = timezone.localtime(period_start)
    local_end = timezone.localtime(period_end)
    end_hour = 24 if local_end.hour == 0 and local_end.date() > local_start.date() else local_end.hour
    return f"{local_start.hour:02d}:00-{end_hour:02d}:00"


def _fmt(value):
    return timezone.localtime(value).strftime("%Y-%m-%d %H:%M")
